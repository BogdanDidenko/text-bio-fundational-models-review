#!/usr/bin/env python3
"""Build analysis tables, deterministic figures, and the HIC 2026 abstract DOCX."""

from __future__ import annotations

import argparse
import collections
import csv
import hashlib
import html
import json
import math
import shutil
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TAXONOMY = ROOT / "data/input_representation_taxonomy_2026-07-11"
DEFAULT_TEMPLATE = ROOT / (
    "data/health_intelligence_conference_2026_template/"
    "health_intelligence_conference_2026_abstract_template_EN.docx"
)
DEFAULT_OUTPUT = ROOT / "data/health_intelligence_conference_2026_abstract_2026-07-11"
MOTIVATION_SOURCE = DEFAULT_OUTPUT / "motivation_synthesis/motivation_synthesis.json"


FAMILY_ORDER = [
    "text_native_token_stream",
    "dense_continuous_carrier",
    "visual_raster_carrier",
    "discrete_biological_symbol_stream",
    "geometric_or_diffusion_state_carrier",
]

FAMILY_META = {
    "text_native_token_stream": {
        "label": "Text-native token streams",
        "definition": "Ordinary tokenizer-visible text",
        "fill": "EAF3FF",
        "chip": "C9DFFF",
    },
    "dense_continuous_carrier": {
        "label": "Dense continuous carriers",
        "definition": "Vectors entering without ordinary tokens",
        "fill": "EAF8F0",
        "chip": "C8EBD7",
    },
    "visual_raster_carrier": {
        "label": "Visual raster carriers",
        "definition": "Pixels, patches, and slide context",
        "fill": "FFF5E6",
        "chip": "FFE0AD",
    },
    "discrete_biological_symbol_stream": {
        "label": "Discrete biological symbol streams",
        "definition": "Native or learned biological token IDs",
        "fill": "FFF0F5",
        "chip": "F7C9DB",
    },
    "geometric_or_diffusion_state_carrier": {
        "label": "Geometric or diffusion-state carriers",
        "definition": "Geometry- or time-indexed model states",
        "fill": "F3EFFF",
        "chip": "DCD2FA",
    },
}

SUBTYPE_ORDER = {
    "text_native_token_stream": [
        "plain_language_prompt_or_question",
        "structured_biological_prompt_or_task_scaffold",
        "serialized_biological_context_or_ordered_profile",
    ],
    "dense_continuous_carrier": [
        "direct_projected_embedding",
        "virtual_token_prefix",
        "connector_mediated_embedding",
        "pooled_or_aggregated_embedding",
    ],
    "visual_raster_carrier": [
        "raw_slide_or_patch_input",
        "patch_context_or_case_level_visual_reasoning",
    ],
    "discrete_biological_symbol_stream": [
        "native_biological_token_stream",
        "multi_track_structural_symbol_stream",
        "learned_quantized_id_or_codebook_token",
    ],
    "geometric_or_diffusion_state_carrier": [
        "noisy_diffusion_state",
        "coordinate_backbone_or_shape_conditioning",
        "symbolic_structural_constraint",
    ],
}

SUBTYPE_LABELS = {
    "plain_language_prompt_or_question": "Plain-language prompts or questions",
    "structured_biological_prompt_or_task_scaffold": "Structured biological prompts/task scaffolds",
    "serialized_biological_context_or_ordered_profile": "Serialized biological contexts/ordered profiles",
    "direct_projected_embedding": "Direct projected embeddings",
    "virtual_token_prefix": "Virtual-token prefixes",
    "connector_mediated_embedding": "Connector-mediated embeddings",
    "pooled_or_aggregated_embedding": "Pooled/aggregated embeddings",
    "raw_slide_or_patch_input": "Raw slide or patch input",
    "patch_context_or_case_level_visual_reasoning": "Patch-context/case-level visual reasoning",
    "native_biological_token_stream": "Native biological token streams",
    "multi_track_structural_symbol_stream": "Multi-track structural symbol streams",
    "learned_quantized_id_or_codebook_token": "Learned quantized IDs/codebook tokens",
    "noisy_diffusion_state": "Noisy diffusion state",
    "coordinate_backbone_or_shape_conditioning": "Coordinate/backbone/shape conditioning",
    "symbolic_structural_constraint": "Symbolic structural constraints",
}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line]


def build_facts(taxonomy_dir: Path) -> dict[str, Any]:
    rows = read_jsonl(taxonomy_dir / "route_annotations.jsonl")
    family_counts = collections.Counter(r["carrier_family"] for r in rows)
    subtype_counts = collections.Counter(r["carrier_subtype"] for r in rows)
    family_records: dict[str, set[str]] = collections.defaultdict(set)
    record_families: dict[str, set[str]] = collections.defaultdict(set)
    configuration_families: dict[str, set[str]] = collections.defaultdict(set)
    record_configurations: dict[str, set[str]] = collections.defaultdict(set)
    for row in rows:
        family_records[row["carrier_family"]].add(row["record_id"])
        record_families[row["record_id"]].add(row["carrier_family"])
        configuration_families[row["configuration_id"]].add(row["carrier_family"])
        record_configurations[row["record_id"]].add(row["configuration_id"])
    agreement = read_json(taxonomy_dir / "agreement_metrics.json")
    registry = read_json(taxonomy_dir / "registry_summary.json")
    values_per_record = collections.Counter(r["record_id"] for r in rows)
    ordered_route_counts = sorted(values_per_record.values())
    median = (
        ordered_route_counts[len(ordered_route_counts) // 2]
        if len(ordered_route_counts) % 2
        else sum(ordered_route_counts[len(ordered_route_counts) // 2 - 1 : len(ordered_route_counts) // 2 + 1]) / 2
    )
    families = []
    for family in FAMILY_ORDER:
        subtypes = []
        for subtype in SUBTYPE_ORDER[family]:
            subtypes.append(
                {
                    "id": subtype,
                    "label": SUBTYPE_LABELS[subtype],
                    "routes": subtype_counts[subtype],
                }
            )
        families.append(
            {
                "id": family,
                "label": FAMILY_META[family]["label"],
                "definition": FAMILY_META[family]["definition"],
                "routes": family_counts[family],
                "records": len(family_records[family]),
                "subtypes": subtypes,
            }
        )
    if sum(x["routes"] for x in families) != len(rows):
        raise ValueError("Family route counts do not sum to accepted routes")
    for family in families:
        if sum(s["routes"] for s in family["subtypes"]) != family["routes"]:
            raise ValueError(f"Subtype counts do not sum for {family['id']}")
    review_flow = {
        "databases": 7,
        "search_rounds": 4,
        "records_identified": 7531,
        "duplicates_removed": 2729,
        "unique_or_new_after_dedup": 4802,
        "without_usable_abstract": 225,
        "title_abstract_screened": 4577,
        "title_abstract_excluded": 4327,
        "full_text_candidates": 250,
        "reports_not_retrieved": 15,
        "reports_retrieved": 235,
        "without_valid_targeted_section_pair": 14,
        "targeted_section_screened": 221,
        "automated_include": 50,
        "automated_exclude": 165,
        "automated_uncertain": 6,
        "manual_resolution_include": 2,
        "manual_resolution_exclude": 4,
        "accepted_records": 52,
        "full_text_excluded_after_resolution": 169,
    }
    cross_configuration_only_records = 0
    for record_id, families_for_record in record_families.items():
        if len(families_for_record) <= 1:
            continue
        if all(
            len(configuration_families[configuration_id]) == 1
            for configuration_id in record_configurations[record_id]
        ):
            cross_configuration_only_records += 1

    rna_carriers = collections.Counter(
        row["carrier_family"]
        for row in rows
        if row["source_modality_normalized"] == "RNA"
    )
    dense_fusion = collections.Counter(
        row["fusion_topology"]
        for row in rows
        if row["carrier_family"] == "dense_continuous_carrier"
    )
    discovery_candidates = read_json(
        taxonomy_dir / "taxonomy_synthesis/open_route_inventory.json"
    )
    motivation = read_json(MOTIVATION_SOURCE)
    configuration_route_counts = collections.Counter(
        row["configuration_id"] for row in rows
    )
    screening_summaries = [
        read_json(ROOT / "data/screening_codex_full_2026-07-06/summary.json"),
        read_json(ROOT / "data/screening_codex_2026-06-10/summary.json"),
        read_json(ROOT / "data/screening_codex_2026-07-06/summary.json"),
    ]
    assert sum(x["total_records"] for x in screening_summaries) == 4577
    assert sum(x["decision_counts"].get("EXCLUDE", 0) for x in screening_summaries) == 4327
    assert sum(x["decision_counts"].get("INCLUDE", 0) for x in screening_summaries) == 155
    assert sum(x["decision_counts"].get("UNCERTAIN", 0) for x in screening_summaries) == 95
    full_summary = read_json(
        ROOT
        / "data/screening_codex_fulltext_docling_graph_direct_clean_both_targets_2026-07-10/summary.json"
    )
    assert full_summary["total_records"] == 221
    context_summary = read_json(
        ROOT
        / "data/fulltext_screening_context_2026-07-10_docling_graph_direct_all235_clean_both_targets/run_metadata.json"
    )
    assert context_summary["source_records_before_filter"] == 235
    facts = {
        "classification_unit": "study -> model -> lifecycle phase -> task/input configuration -> input route",
        "route_definition": "one source object followed through its transformation chain to the model-visible carrier consumed by the generative model",
        "records": len(record_families),
        "primary_studies": registry.get("primary_study_count", 51),
        "sensitivity_studies": registry.get("sensitivity_study_count_if_omnina_linked", 50),
        "models": len({r["model_id"] for r in rows}),
        "configurations": len({r["configuration_id"] for r in rows}),
        "routes": len(rows),
        "routes_per_record": {
            "min": min(values_per_record.values()),
            "median": median,
            "max": max(values_per_record.values()),
        },
        "single_family_records": sum(len(v) == 1 for v in record_families.values()),
        "multi_family_records": sum(len(v) > 1 for v in record_families.values()),
        "single_family_configurations": sum(len(v) == 1 for v in configuration_families.values()),
        "multi_family_configurations": sum(len(v) > 1 for v in configuration_families.values()),
        "configuration_route_cardinality": {
            "one_route": sum(value == 1 for value in configuration_route_counts.values()),
            "multiple_routes": sum(value > 1 for value in configuration_route_counts.values()),
            "maximum": max(configuration_route_counts.values()),
        },
        "cross_configuration_only_multi_family_records": cross_configuration_only_records,
        "families": families,
        "text_roles": dict(collections.Counter(r["text_role"] for r in rows).most_common()),
        "fusion_topologies": dict(collections.Counter(r["fusion_topology"] for r in rows).most_common()),
        "lifecycle_phases": dict(collections.Counter(r["lifecycle_phase"] for r in rows).most_common()),
        "input_status": dict(collections.Counter(r["input_status"] for r in rows).most_common()),
        "rna_carrier_counts": dict(rna_carriers),
        "dense_fusion_counts": dict(dense_fusion),
        "open_discovery_candidates": len(discovery_candidates),
        "grounding": {
            "validated": sum(bool(r["final_grounding_valid"]) for r in rows),
            "exact_canonical_markdown": sum(
                bool(r["quote_verified_in_canonical_markdown"]) for r in rows
            ),
            "native_docling_item_fallback": sum(
                not bool(r["quote_verified_in_canonical_markdown"]) for r in rows
            ),
            "picture_only": sum(bool(r["picture_only_provenance"]) for r in rows),
        },
        "motivation": {
            "verified_claims": motivation["run_metadata"]["verified_claim_count"],
            "records": motivation["run_metadata"]["records_with_verified_claims"],
        },
        "agreement": agreement,
        "review_flow": review_flow,
        "source_artifacts": {
            "route_annotations": str(taxonomy_dir / "route_annotations.jsonl"),
            "route_annotations_sha256": sha256(taxonomy_dir / "route_annotations.jsonl"),
            "taxonomy_codebook": str(taxonomy_dir / "taxonomy_codebook.md"),
            "taxonomy_codebook_sha256": sha256(taxonomy_dir / "taxonomy_codebook.md"),
            "agreement_metrics": str(taxonomy_dir / "agreement_metrics.json"),
            "agreement_metrics_sha256": sha256(taxonomy_dir / "agreement_metrics.json"),
        },
    }
    assert facts["records"] == 52
    assert facts["routes"] == 489
    assert facts["models"] == 111
    assert facts["configurations"] == 376
    assert facts["multi_family_records"] == 37
    assert facts["single_family_configurations"] == 332
    assert facts["multi_family_configurations"] == 44
    assert facts["configuration_route_cardinality"] == {
        "one_route": 299,
        "multiple_routes": 77,
        "maximum": 7,
    }
    assert facts["cross_configuration_only_multi_family_records"] == 15
    assert facts["open_discovery_candidates"] == 583
    assert facts["grounding"]["validated"] == 489
    assert facts["grounding"]["picture_only"] == 0
    assert facts["motivation"] == {"verified_claims": 174, "records": 52}
    return facts


def write_analysis_tables(facts: dict[str, Any], output: Path) -> None:
    analysis = output / "analysis"
    analysis.mkdir(parents=True, exist_ok=True)
    (analysis / "manuscript_fact_table.json").write_text(
        json.dumps(facts, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    with (analysis / "taxonomy_frequencies.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["family_id", "family", "subtype_id", "subtype", "routes", "family_records"])
        for family in facts["families"]:
            for subtype in family["subtypes"]:
                writer.writerow(
                    [
                        family["id"],
                        family["label"],
                        subtype["id"],
                        subtype["label"],
                        subtype["routes"],
                        family["records"],
                    ]
                )
    lines = [
        "# Conference Manuscript Fact Table",
        "",
        f"- Records: **{facts['records']}**",
        f"- Primary studies: **{facts['primary_studies']}**",
        f"- Models: **{facts['models']}**",
        f"- Configurations: **{facts['configurations']}**",
        f"- Grounded routes: **{facts['routes']}**",
        f"- Multi-family records: **{facts['multi_family_records']}/{facts['records']}**",
        f"- Routes per record: **{facts['routes_per_record']['min']}–{facts['routes_per_record']['max']}**, median **{facts['routes_per_record']['median']}**",
        "",
        "| Family | Routes | Records |",
        "|---|---:|---:|",
    ]
    for family in facts["families"]:
        lines.append(f"| {family['label']} | {family['routes']} | {family['records']} |")
    agreement = facts["agreement"]
    lines.extend(
        [
            "",
            f"- Minimum route-detection Jaccard: **{agreement['minimum_pairwise_jaccard']:.3f}**",
            f"- Carrier-family exact agreement: **{agreement['carrier_family_exact_agreement']:.3f}**",
            f"- Carrier-family Krippendorff alpha: **{agreement['carrier_family_krippendorff_alpha']:.3f}**",
        ]
    )
    (analysis / "manuscript_fact_table.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    prisma_labels = {
        "records_identified": "Records identified across seven databases and four search rounds",
        "duplicates_removed": "Duplicates removed",
        "unique_or_new_after_dedup": "Unique or new records after deduplication",
        "without_usable_abstract": "Records without usable abstracts",
        "title_abstract_screened": "Records screened by title and abstract",
        "title_abstract_excluded": "Records excluded at title/abstract screening",
        "full_text_candidates": "Full-text candidates",
        "reports_not_retrieved": "Reports not retrieved",
        "reports_retrieved": "Reports retrieved",
        "without_valid_targeted_section_pair": "Reports without a valid targeted-section pair",
        "targeted_section_screened": "Reports screened using complete selected sections",
        "automated_include": "Automated full-section INCLUDE",
        "automated_exclude": "Automated full-section EXCLUDE",
        "automated_uncertain": "Automated full-section UNCERTAIN",
        "manual_resolution_include": "Manual resolution to INCLUDE",
        "manual_resolution_exclude": "Manual resolution to EXCLUDE",
        "accepted_records": "Accepted records",
        "full_text_excluded_after_resolution": "Full-section exclusions after resolution",
    }
    with (analysis / "prisma_fact_table.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["fact_id", "description", "count"])
        for fact_id, description in prisma_labels.items():
            writer.writerow([fact_id, description, facts["review_flow"][fact_id]])
    prisma_lines = ["# Verified PRISMA fact table", "", "| Stage | Count |", "|---|---:|"]
    for fact_id, description in prisma_labels.items():
        prisma_lines.append(f"| {description} | {facts['review_flow'][fact_id]:,} |")
    (analysis / "prisma_fact_table.md").write_text("\n".join(prisma_lines) + "\n", encoding="utf-8")


def svg_text(x: float, y: float, text: str, cls: str, anchor: str = "start") -> str:
    return f'<text x="{x}" y="{y}" text-anchor="{anchor}" class="{cls}">{html.escape(text)}</text>'


def taxonomy_svg(facts: dict[str, Any]) -> str:
    width, height = 1800, 600
    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#ffffff"/>',
        "<style>",
        ".route{font:700 27px Arial,sans-serif;fill:#152033}.helper{font:17px Arial,sans-serif;fill:#475569}",
        ".header{font:700 25px Arial,sans-serif;fill:#152033}.family{font:700 24px Arial,sans-serif;fill:#152033}",
        ".definition{font:20px Arial,sans-serif;fill:#475569}.leaf{font:19px Arial,sans-serif;fill:#152033}.leafcount{font:700 18px Arial,sans-serif;fill:#334155}",
        ".total{font:700 23px Arial,sans-serif;fill:#152033}.footer{font:700 18px Arial,sans-serif;fill:#334155}.footlight{font:18px Arial,sans-serif;fill:#475569}",
        "</style>",
        '<defs><marker id="arrow" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#475569"/></marker></defs>',
    ]
    steps = [
        (46, 18, 300, "Source object", "text · sequence · omics · image · structure"),
        (410, 18, 330, "Transformation chain", "serialize · tokenize · project · pool · encode"),
        (804, 18, 340, "Model-visible carrier", "primary taxonomy axis"),
        (1208, 18, 546, "Generative backbone", "pretraining · adaptation · evaluation · inference"),
    ]
    for idx, (x, y, w, title, helper) in enumerate(steps):
        fill = "#fff7ed" if idx == 2 else "#f8fafc"
        stroke = "#d97706" if idx == 2 else "#94a3b8"
        parts.append(f'<rect x="{x}" y="{y}" width="{w}" height="70" rx="7" fill="{fill}" stroke="{stroke}" stroke-width="2"/>')
        parts.append(svg_text(x + w / 2, y + 30, title, "route", "middle"))
        parts.append(svg_text(x + w / 2, y + 55, helper, "helper", "middle"))
        if idx < len(steps) - 1:
            nx = steps[idx + 1][0]
            parts.append(f'<path d="M{x + w} {y + 35} H{nx - 18}" stroke="#475569" stroke-width="3" marker-end="url(#arrow)"/>')
    parts.extend(
        [
            svg_text(46, 118, "Carrier family", "header"),
            svg_text(540, 118, "Evidence-grounded subtypes", "header"),
            svg_text(1698, 118, "Routes · records", "header", "end"),
        ]
    )
    row_y = 128
    row_h = 75
    gap = 5
    for family in facts["families"]:
        meta = FAMILY_META[family["id"]]
        parts.append(f'<rect x="36" y="{row_y}" width="1728" height="{row_h}" rx="7" fill="#{meta["fill"]}"/>')
        parts.append(svg_text(56, row_y + 30, family["label"], "family"))
        parts.append(svg_text(56, row_y + 56, family["definition"], "definition"))
        subtype_x = 530
        subtype_space = 1030
        n = len(family["subtypes"])
        chip_gap = 12
        chip_w = (subtype_space - chip_gap * (n - 1)) / n
        for subtype in family["subtypes"]:
            parts.append(
                f'<rect x="{subtype_x}" y="{row_y + 8}" width="{chip_w}" height="59" rx="6" '
                f'fill="#{meta["chip"]}" aria-label="{html.escape(subtype["label"])}"/>'
            )
            label = subtype["label"]
            wrap_at = {2: 44, 3: 31, 4: 24}[n]
            if len(label) > wrap_at:
                split = label.rfind(" ", 0, wrap_at + 1)
                split = split if split > 12 else wrap_at
                parts.append(svg_text(subtype_x + 14, row_y + 28, label[:split], "leaf"))
                parts.append(svg_text(subtype_x + 14, row_y + 50, label[split + 1 :], "leaf"))
                parts.append(svg_text(subtype_x + chip_w - 14, row_y + 58, str(subtype["routes"]), "leafcount", "end"))
            else:
                parts.append(svg_text(subtype_x + 14, row_y + 33, label, "leaf"))
                parts.append(svg_text(subtype_x + 14, row_y + 58, str(subtype["routes"]), "leafcount"))
            subtype_x += chip_w + chip_gap
        parts.append(svg_text(1698, row_y + 43, f"{family['routes']} · {family['records']}", "total", "end"))
        row_y += row_h + gap
    footer_y = 535
    parts.append(f'<rect x="36" y="{footer_y}" width="1728" height="48" rx="7" fill="#f8fafc" stroke="#cbd5e1"/>')
    parts.append(svg_text(56, footer_y + 31, "Orthogonal annotations: source modality · text role · lifecycle phase · fusion topology", "footlight"))
    summary = (
        f"{facts['records']} records · {facts['models']} models · {facts['configurations']} configurations · "
        f"{facts['routes']} routes · {facts['multi_family_records']} multi-family records"
    )
    parts.append(svg_text(1740, footer_y + 31, summary, "footer", "end"))
    parts.append("</svg>")
    return "\n".join(parts)


def composite_taxonomy_svg(facts: dict[str, Any]) -> str:
    """Compact one-page figure: taxonomy, empirical findings, and evidence audit."""
    width, height = 1800, 640
    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#ffffff"/>',
        "<style>",
        ".route{font:700 25px Arial,sans-serif;fill:#172033}.helper{font:20px Arial,sans-serif;fill:#475569}",
        ".panel{font:700 24px Arial,sans-serif;fill:#172033}.header{font:700 22px Arial,sans-serif;fill:#172033}",
        ".family{font:700 23px Arial,sans-serif;fill:#172033}.definition{font:20px Arial,sans-serif;fill:#475569}",
        ".leaf{font:21px Arial,sans-serif;fill:#172033}.count{font:700 22px Arial,sans-serif;fill:#172033}",
        ".small{font:21px Arial,sans-serif;fill:#334155}.smallb{font:700 22px Arial,sans-serif;fill:#172033}",
        ".metric{font:700 26px Arial,sans-serif;fill:#172033}.footer{font:700 20px Arial,sans-serif;fill:#334155}",
        "</style>",
        '<defs><marker id="arrow2" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#475569"/></marker></defs>',
    ]

    steps = [
        (24, 8, 300, "Source object", "text · sequence · omics"),
        (380, 8, 330, "Transformation chain", "tokenize · project · encode"),
        (766, 8, 340, "Model-visible carrier", "primary taxonomy axis"),
        (1162, 8, 614, "Generative backbone", "pretrain · adapt · evaluate · infer"),
    ]
    for idx, (x, y, w, title, helper) in enumerate(steps):
        fill = "#fff7ed" if idx == 2 else "#f8fafc"
        stroke = "#b45309" if idx == 2 else "#94a3b8"
        parts.append(f'<rect x="{x}" y="{y}" width="{w}" height="52" rx="6" fill="{fill}" stroke="{stroke}" stroke-width="2"/>')
        parts.append(svg_text(x + w / 2, y + 23, title, "route", "middle"))
        parts.append(svg_text(x + w / 2, y + 44, helper, "helper", "middle"))
        if idx < len(steps) - 1:
            nx = steps[idx + 1][0]
            parts.append(f'<path d="M{x + w} {y + 26} H{nx - 18}" stroke="#475569" stroke-width="3" marker-end="url(#arrow2)"/>')

    parts.extend(
        [
            svg_text(24, 84, "Carrier family", "header"),
            svg_text(430, 84, "Evidence-grounded subtypes · routes", "header"),
            svg_text(1772, 84, "Routes · records", "header", "end"),
        ]
    )
    display_labels = {
        "plain_language_prompt_or_question": "Plain-language prompts/questions",
        "structured_biological_prompt_or_task_scaffold": "Structured bio prompts/scaffolds",
        "serialized_biological_context_or_ordered_profile": "Serialized bio contexts/profiles",
        "direct_projected_embedding": "Direct projection",
        "virtual_token_prefix": "Virtual-token prefix",
        "connector_mediated_embedding": "Connector-mediated",
        "pooled_or_aggregated_embedding": "Pooled/aggregated",
        "raw_slide_or_patch_input": "Raw slides/patches",
        "patch_context_or_case_level_visual_reasoning": "Contextual patch/case reasoning",
        "native_biological_token_stream": "Native bio token streams",
        "multi_track_structural_symbol_stream": "Multi-track structural symbols",
        "learned_quantized_id_or_codebook_token": "Learned quantized/codebook IDs",
        "noisy_diffusion_state": "Noisy diffusion states",
        "coordinate_backbone_or_shape_conditioning": "Coordinate/backbone/shape",
        "symbolic_structural_constraint": "Symbolic structural constraints",
    }
    family_display_labels = {
        "text_native_token_stream": "Text-native token streams",
        "dense_continuous_carrier": "Dense continuous carriers",
        "visual_raster_carrier": "Visual raster carriers",
        "discrete_biological_symbol_stream": "Discrete biological symbols",
        "geometric_or_diffusion_state_carrier": "Geometry/diffusion states",
    }
    row_y, row_h, gap = 92, 48, 4
    for family in facts["families"]:
        meta = FAMILY_META[family["id"]]
        parts.append(
            f'<rect x="24" y="{row_y}" width="1752" height="{row_h}" rx="7" '
            f'fill="#{meta["fill"]}" stroke="#d5dbe3" aria-label="{html.escape(family["label"])}"/>'
        )
        parts.append(svg_text(42, row_y + 21, family_display_labels[family["id"]], "family"))
        parts.append(svg_text(42, row_y + 41, meta["definition"], "definition"))
        subtype_x, subtype_space = 420, 1190
        n = len(family["subtypes"])
        chip_gap = 8
        chip_w = (subtype_space - chip_gap * (n - 1)) / n
        for subtype in family["subtypes"]:
            parts.append(
                f'<rect x="{subtype_x}" y="{row_y + 5}" width="{chip_w}" height="38" rx="5" '
                f'fill="#{meta["chip"]}" aria-label="{html.escape(subtype["label"])}"/>'
            )
            parts.append(svg_text(subtype_x + 10, row_y + 30, display_labels[subtype["id"]], "leaf"))
            parts.append(svg_text(subtype_x + chip_w - 9, row_y + 30, str(subtype["routes"]), "count", "end"))
            subtype_x += chip_w + chip_gap
        parts.append(svg_text(1752, row_y + 31, f'{family["routes"]} · {family["records"]}', "metric", "end"))
        row_y += row_h + gap

    panel_y, panel_h = 356, 190
    panels = [(24, 515), (551, 590), (1153, 623)]
    for x, w in panels:
        parts.append(f'<rect x="{x}" y="{panel_y}" width="{w}" height="{panel_h}" rx="8" fill="#f8fafc" stroke="#cbd5e1" stroke-width="2"/>')

    x = 44
    parts.append(svg_text(x, panel_y + 27, "B  Multimodality by denominator", "panel"))
    for label, numerator, denominator, color, y in [
        ("Records", facts["multi_family_records"], facts["records"], "#4f78a8", panel_y + 57),
        ("Task-input settings", facts["multi_family_configurations"], facts["configurations"], "#e07b39", panel_y + 101),
    ]:
        pct = numerator / denominator
        parts.append(svg_text(x, y, label, "smallb"))
        parts.append(f'<rect x="{x}" y="{y + 7}" width="430" height="14" rx="4" fill="#e2e8f0"/>')
        parts.append(f'<rect x="{x}" y="{y + 7}" width="{430 * pct:.1f}" height="14" rx="4" fill="{color}"/>')
        parts.append(svg_text(x + 440, y, f"{numerator}/{denominator} ({100 * pct:.1f}%)", "metric", "end"))
    parts.append(svg_text(x, panel_y + 153, f'{facts["single_family_configurations"]}/{facts["configurations"]} settings expose one family', "smallb"))
    parts.append(svg_text(x, panel_y + 177, f'{facts["cross_configuration_only_multi_family_records"]} papers: multi-family only across tasks/phases', "small"))

    x = 571
    parts.append(svg_text(x, panel_y + 27, "C  Text has multiple roles", "panel"))
    text_role_labels = [
        ("instruction/query", "instruction_or_query"),
        ("bio payload", "biological_payload"),
        ("alignment", "paired_alignment_supervision"),
        ("no text", "no_text_on_this_route"),
        ("metadata/context", "metadata_or_context"),
        ("task/modality cue", "modality_or_task_selector"),
        ("generated output", "generated_output"),
        ("semantic label", "semantic_annotation"),
        ("unclear", "unclear"),
    ]
    for idx, (label, key) in enumerate(text_role_labels):
        col, row = idx % 3, idx // 3
        bx = x + col * 190
        by = panel_y + 45 + row * 46
        value = facts["text_roles"][key]
        parts.append(svg_text(bx, by + 16, label, "small"))
        parts.append(svg_text(bx, by + 38, str(value), "count"))

    x = 1173
    parts.append(svg_text(x, panel_y + 27, "D  Orthogonal design choices", "panel"))
    parts.append(svg_text(x, panel_y + 55, "Same source · RNA routes", "smallb"))
    rna_values = [
        ("text-native", facts["rna_carrier_counts"].get("text_native_token_stream", 0), "#78a7d8"),
        ("dense", facts["rna_carrier_counts"].get("dense_continuous_carrier", 0), "#74b58b"),
        ("discrete", facts["rna_carrier_counts"].get("discrete_biological_symbol_stream", 0), "#d989aa"),
    ]
    bx, by = x, panel_y + 66
    for idx, (label, value, color) in enumerate(rna_values):
        chip_x = bx + idx * 190
        parts.append(f'<rect x="{chip_x}" y="{by}" width="178" height="27" rx="5" fill="{color}"/>')
        parts.append(svg_text(chip_x + 89, by + 21, f"{label} {value}", "smallb", "middle"))
    parts.append(svg_text(x, panel_y + 119, "Same carrier · dense fusion", "smallb"))
    dense_items = [
        ("cross-attention", "cross_attention"),
        ("placeholder", "placeholder_replacement"),
        ("concatenation", "concatenation"),
        ("prefix", "prefix"),
    ]
    for idx, (label, key) in enumerate(dense_items):
        col, row = idx % 2, idx // 2
        tx, ty = x + col * 286, panel_y + 148 + row * 30
        value = facts["dense_fusion_counts"][key]
        parts.append(svg_text(tx, ty, label, "small"))
        parts.append(svg_text(tx + 260, ty, str(value), "count", "end"))

    footer_y = 556
    parts.append(f'<rect x="24" y="{footer_y}" width="1752" height="68" rx="8" fill="#eef2f6" stroke="#94a3b8" stroke-width="2"/>')
    prisma = facts["review_flow"]
    agreement = facts["agreement"]
    footer_lines = [
        f'PRISMA-ScR: 7 databases · 4 rounds · {prisma["records_identified"]:,} identified → {prisma["title_abstract_screened"]:,} title/abstract → {prisma["targeted_section_screened"]} sections → {facts["records"]} records / {facts["primary_studies"]} studies',
        f'Corpus: {facts["models"]} models · {facts["configurations"]} settings · Extraction: {facts["open_discovery_candidates"]} → {facts["routes"]} routes · 2–34/record (median 8) · Dense: {agreement["dense_candidate_count"]:,} → +{agreement["dense_only_accepted_candidate_count"]} · Provenance: {facts["grounding"]["validated"]}/{facts["routes"]}',
        f'Agreement: route Jaccard ≥ {agreement["minimum_pairwise_jaccard"]:.3f} · family exact agreement {agreement["carrier_family_exact_agreement"]:.3f} · Krippendorff α {agreement["carrier_family_krippendorff_alpha"]:.3f}',
    ]
    for idx, line in enumerate(footer_lines):
        parts.append(svg_text(44, footer_y + 20 + idx * 21, line, "footer"))
    parts.append("</svg>")
    return "\n".join(parts)


def workflow_svg(facts: dict[str, Any]) -> str:
    flow = facts["review_flow"]
    width, height = 1800, 185
    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#ffffff"/>',
        "<style>.h{font:700 23px Arial,sans-serif;fill:#152033}.n{font:700 18px Arial,sans-serif;fill:#152033}.s{font:17px Arial,sans-serif;fill:#475569}.tiny{font:16px Arial,sans-serif;fill:#475569}</style>",
        '<defs><marker id="a" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#64748b"/></marker></defs>',
        svg_text(38, 24, "Evidence assembly and repeated computational screening", "h"),
    ]
    funnel = [
        (40, "Search", f"7 databases · 4 rounds · {flow['records_identified']:,}"),
        (250, "Dedup + abstracts", f"{flow['title_abstract_screened']:,} screened"),
        (460, "Title/abstract", f"{flow['full_text_candidates']} candidates"),
        (670, "Full text", f"{flow['reports_retrieved']} retrieved"),
        (880, "Docling + Graph", f"{flow['targeted_section_screened']} section inputs"),
        (1090, "Full-section", f"{flow['accepted_records']} included"),
    ]
    for idx, (x, title, count) in enumerate(funnel):
        w = 190
        parts.append(f'<rect x="{x}" y="36" width="{w}" height="56" rx="6" fill="#f8fafc" stroke="#94a3b8" stroke-width="1.6"/>')
        parts.append(svg_text(x + w / 2, 59, title, "n", "middle"))
        parts.append(svg_text(x + w / 2, 81, count, "s", "middle"))
        if idx < len(funnel) - 1:
            nx = funnel[idx + 1][0]
            parts.append(f'<path d="M{x + w} 64 H{nx - 13}" stroke="#64748b" stroke-width="2.2" marker-end="url(#a)"/>')
    parts.append(f'<rect x="1330" y="30" width="430" height="72" rx="7" fill="#eef6ff" stroke="#6b8fbd" stroke-width="1.8"/>')
    parts.append(svg_text(1545, 54, "Taxonomy evidence corpus", "n", "middle"))
    parts.append(svg_text(1545, 76, f"{facts['primary_studies']} studies · {facts['models']} models · {facts['routes']} routes", "s", "middle"))
    parts.append(svg_text(1545, 95, "complete VLM-enriched Docling profiles", "tiny", "middle"))
    parts.append('<path d="M1295 64 H1317" stroke="#64748b" stroke-width="2.2" marker-end="url(#a)"/>')
    parts.append(f'<rect x="40" y="112" width="815" height="55" rx="7" fill="#fff7ed" stroke="#d8a45f"/>')
    parts.append(svg_text(60, 136, "Pass 1 input:", "n"))
    parts.append(svg_text(185, 136, "title + abstract", "s"))
    parts.append(svg_text(60, 158, "scope reviewer + architecture reviewer → Python gate → adjudicator", "s"))
    parts.append(f'<rect x="875" y="112" width="885" height="55" rx="7" fill="#edf8f1" stroke="#70a585"/>')
    parts.append(svg_text(895, 136, "Pass 2 input:", "n"))
    parts.append(svg_text(1020, 136, "title + abstract + complete selected sections", "s"))
    parts.append(svg_text(895, 158, "PDF → Docling → Docling Graph provenance → repeated role pipeline → manual resolution", "s"))
    parts.append("</svg>")
    return "\n".join(parts)


def render_svg(svg_path: Path, png_path: Path, width: int) -> None:
    candidates = [shutil.which("rsvg-convert"), "/usr/local/bin/rsvg-convert"]
    exe = next((x for x in candidates if x and Path(x).exists()), None)
    if not exe:
        raise RuntimeError("rsvg-convert is required to render deterministic SVG previews")
    cmd = [exe, "-w", str(width), "-o", str(png_path), str(svg_path)]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0 and Path("/usr/bin/arch").exists():
        result = subprocess.run(["arch", "-x86_64", *cmd], capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(result.stderr)


def validate_figure_text(
    facts: dict[str, Any], taxonomy_path: Path, workflow_path: Path, variant: str = "taxonomy"
) -> dict[str, Any]:
    taxonomy = taxonomy_path.read_text(encoding="utf-8")
    workflow = workflow_path.read_text(encoding="utf-8")
    taxonomy_expected = [
        "Source object",
        "Transformation chain",
        "Model-visible carrier",
        "Generative backbone",
        *[family["label"] for family in facts["families"]],
        *[subtype["label"] for family in facts["families"] for subtype in family["subtypes"]],
        *[str(family["routes"]) for family in facts["families"]],
        *[str(family["records"]) for family in facts["families"]],
        str(facts["records"]),
        str(facts["models"]),
        str(facts["configurations"]),
        str(facts["routes"]),
    ]
    if variant in {"composite-v2", "compact-v3"}:
        taxonomy_expected.extend(
            [
                "Multimodality by denominator",
                "Text has multiple roles",
                "Orthogonal design choices",
                f'{facts["multi_family_records"]}/{facts["records"]}',
                f'{facts["multi_family_configurations"]}/{facts["configurations"]}',
                str(facts["cross_configuration_only_multi_family_records"]),
                str(facts["text_roles"]["instruction_or_query"]),
                str(facts["dense_fusion_counts"]["cross_attention"]),
                f'{facts["grounding"]["validated"]}/{facts["routes"]}',
            ]
        )
    workflow_expected = [
        f"{facts['review_flow']['records_identified']:,}",
        f"{facts['review_flow']['title_abstract_screened']:,}",
        str(facts["review_flow"]["full_text_candidates"]),
        str(facts["review_flow"]["reports_retrieved"]),
        str(facts["review_flow"]["targeted_section_screened"]),
        str(facts["review_flow"]["accepted_records"]),
        "title + abstract",
        "complete selected sections",
        "Docling Graph provenance",
    ]
    missing_taxonomy = sorted({item for item in taxonomy_expected if html.escape(item) not in taxonomy})
    missing_workflow = sorted({item for item in workflow_expected if html.escape(item) not in workflow})
    report = {
        "taxonomy_expected_text_items": len(taxonomy_expected),
        "workflow_expected_text_items": len(workflow_expected),
        "missing_taxonomy_items": missing_taxonomy,
        "missing_workflow_items": missing_workflow,
        "passed": not missing_taxonomy and not missing_workflow,
    }
    if not report["passed"]:
        raise ValueError(f"Figure fact-contract validation failed: {report}")
    return report


def format_body(paragraph, *, heading=False, caption=False, compact=False, left=False) -> None:
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT if heading or left else WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(1 if heading else 0)
    fmt.space_after = Pt(1)
    fmt.line_spacing = Pt(8.1 if compact else (8.3 if caption else 9.6))
    for run in paragraph.runs:
        run.font.name = "Arial" if heading else "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
        run.font.size = Pt(7.4 if compact else (7.3 if caption else 8.6))
        run.font.bold = heading
        if caption:
            run.font.italic = True
            run.font.color.rgb = RGBColor(68, 68, 68)


def add_section(
    cell, heading: str, text: str, *, compact: bool = False, left: bool = False
) -> None:
    p = cell.add_paragraph()
    p.add_run(heading)
    format_body(p, heading=True)
    p = cell.add_paragraph()
    p.add_run(text)
    format_body(p, compact=compact, left=left)


def add_references(cell, references: list[str]) -> None:
    p = cell.add_paragraph()
    p.add_run("REFERENCES")
    format_body(p, heading=True)
    for reference in references:
        p = cell.add_paragraph()
        p.add_run(reference)
        format_body(p, compact=True, left=True)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.right_indent = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def add_continuation(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.add_run(text)
    format_body(p)


def build_docx(template: Path, output: Path, figures: Path, content: dict[str, Any]) -> Path:
    final_docx = output / "health_intelligence_conference_2026_taxonomy_abstract_EN.docx"
    shutil.copy2(template, final_docx)
    doc = Document(final_docx)
    p = doc.paragraphs
    p[0].text = content["title"]
    p[1].text = content.get("authors", "[Author names]")
    p[2].text = content.get("affiliations", "[Affiliations]")
    p[3].text = content.get("correspondence", "*Correspondence: [email]")
    p[4].text = content.get("received", "Received [DD Month YYYY]")
    p[7].text = "Motivation: " + content["abstract_motivation"]
    p[8].text = "Results: " + content["abstract_results"]
    p[9].text = "Availability and implementation: " + content["availability"]
    for idx in (0, 1, 2, 3, 4, 7, 8, 9):
        for run in p[idx].runs:
            if idx == 0:
                run.font.name = "Arial"
                run.font.size = Pt(12.5)
                run.font.bold = True
            elif idx == 1:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.font.bold = True
            else:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.5)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
        p[idx].paragraph_format.space_before = Pt(0)
        p[idx].paragraph_format.space_after = Pt(0 if idx not in (0, 4) else 2)
        if idx in (1, 2, 3, 4):
            p[idx].paragraph_format.line_spacing = Pt(9.5)
        elif idx in (7, 8, 9):
            p[idx].paragraph_format.line_spacing = Pt(9.4)
    author_marker = p[1].add_run("*")
    author_marker.font.name = "Arial"
    author_marker._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    author_marker.font.size = Pt(7)
    author_marker.font.bold = False
    author_marker.font.superscript = True
    for run in p[2].runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(68, 68, 68)
    for run in p[3].runs:
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(119, 119, 119)
    for run in p[4].runs:
        run.font.size = Pt(7.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(119, 119, 119)
    for idx in (10, 5):
        p[idx]._element.getparent().remove(p[idx]._element)
    body_table = doc.tables[0]
    table_el = body_table._tbl
    body = table_el.getparent()

    def insert_before_body_table(paragraph) -> None:
        paragraph._p.getparent().remove(paragraph._p)
        body.insert(body.index(table_el), paragraph._p)

    figure_p = doc.add_paragraph()
    insert_before_body_table(figure_p)
    figure_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_p.paragraph_format.space_before = Pt(2)
    figure_p.paragraph_format.space_after = Pt(0)
    figure_p.add_run().add_picture(str(figures / "figure_1_taxonomy.png"), width=Inches(6.95))
    cap = doc.add_paragraph()
    insert_before_body_table(cap)
    cap.add_run(content["figure_1_caption"])
    format_body(cap, caption=True)
    if content.get("include_workflow_figure", True):
        workflow_p = doc.add_paragraph()
        insert_before_body_table(workflow_p)
        workflow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        workflow_p.paragraph_format.space_before = Pt(0)
        workflow_p.paragraph_format.space_after = Pt(0)
        workflow_p.add_run().add_picture(str(figures / "figure_2_workflow.png"), width=Inches(6.95))
        workflow_cap = doc.add_paragraph()
        insert_before_body_table(workflow_cap)
        workflow_cap.add_run(content["figure_2_caption"])
        format_body(workflow_cap, caption=True)
    for cell in body_table.rows[0].cells:
        for nested_table in list(cell.tables):
            nested_table._element.getparent().remove(nested_table._element)
        for paragraph in list(cell.paragraphs):
            paragraph._element.getparent().remove(paragraph._element)
    left, right = body_table.rows[0].cells
    add_section(left, "1  INTRODUCTION", content["introduction"])
    add_section(left, "2  METHODS", content["methods"])
    if content.get("results_left") and content.get("results_right"):
        add_section(left, "3  RESULTS", content["results_left"])
        add_continuation(right, content["results_right"])
    else:
        add_section(right, "3  RESULTS", content["results"])
    add_section(right, "4  DISCUSSION AND CONCLUSION", content["discussion"])
    if content.get("include_body_availability", True):
        add_section(right, "AVAILABILITY", content["availability"], left=True)
    add_section(right, "FUNDING", content["funding"], left=True)
    add_references(right, content["references"])
    doc.save(final_docx)
    return final_docx


def write_manuscript_markdown(content: dict[str, Any], output: Path) -> Path:
    results_text = content.get("results") or (
        content.get("results_left", "") + " " + content.get("results_right", "")
    ).strip()
    lines = [
        f"# {content['title']}",
        "",
        f"**Authors:** {content['authors']}  ",
        f"**Affiliations:** {content['affiliations']}  ",
        f"**Correspondence:** {content['correspondence']}",
        "",
        "## Abstract",
        "",
        f"**Motivation:** {content['abstract_motivation']}",
        "",
        f"**Results:** {content['abstract_results']}",
        "",
        f"**Availability and implementation:** {content['availability']}",
        "",
        "## 1 Introduction",
        "",
        content["introduction"],
        "",
        "## 2 Methods",
        "",
        content["methods"],
        "",
        "## 3 Results",
        "",
        results_text,
        "",
        "## 4 Discussion and conclusion",
        "",
        content["discussion"],
        "",
        "## Funding",
        "",
        content["funding"],
        "",
        "## References",
        "",
    ]
    if content.get("include_body_availability", True):
        funding_index = lines.index("## Funding")
        lines[funding_index:funding_index] = [
            "## Availability",
            "",
            content["availability"],
            "",
        ]
    for reference in content["references"]:
        lines.extend([reference, ""])
    path = output / "manuscript_draft.md"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--taxonomy-dir", type=Path, default=DEFAULT_TAXONOMY)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--content", type=Path)
    parser.add_argument(
        "--figure-variant",
        choices=("taxonomy", "composite-v2", "compact-v3"),
        default="taxonomy",
    )
    parser.add_argument("--build-docx", action="store_true")
    args = parser.parse_args()
    taxonomy_dir = args.taxonomy_dir.resolve()
    output = args.output_dir.resolve()
    output.mkdir(parents=True, exist_ok=True)
    facts = build_facts(taxonomy_dir)
    write_analysis_tables(facts, output)
    figures = output / "figures" / "deterministic"
    figures.mkdir(parents=True, exist_ok=True)
    contract = {
        "route_definition": facts["route_definition"],
        "families": facts["families"],
        "orthogonal_dimensions": ["source modality", "text role", "lifecycle phase", "fusion topology"],
        "totals": {
            "records": facts["records"],
            "studies": facts["primary_studies"],
            "models": facts["models"],
            "configurations": facts["configurations"],
            "routes": facts["routes"],
            "multi_family_records": facts["multi_family_records"],
        },
        "agreement": facts["agreement"],
        "review_flow": facts["review_flow"],
        "text_roles": facts["text_roles"],
        "lifecycle_phases": facts["lifecycle_phases"],
        "configuration_route_cardinality": facts["configuration_route_cardinality"],
        "rna_carrier_counts": facts["rna_carrier_counts"],
        "dense_fusion_counts": facts["dense_fusion_counts"],
        "cross_configuration_only_multi_family_records": facts[
            "cross_configuration_only_multi_family_records"
        ],
        "open_discovery_candidates": facts["open_discovery_candidates"],
        "grounding": facts["grounding"],
        "motivation": facts["motivation"],
    }
    (output / "analysis" / "figure_fact_contract.json").write_text(
        json.dumps(contract, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    f1 = figures / "figure_1_taxonomy.svg"
    f2 = figures / "figure_2_workflow.svg"
    figure_svg = (
        composite_taxonomy_svg(facts)
        if args.figure_variant in {"composite-v2", "compact-v3"}
        else taxonomy_svg(facts)
    )
    f1.write_text(figure_svg, encoding="utf-8")
    f2.write_text(workflow_svg(facts), encoding="utf-8")
    render_svg(f1, figures / "figure_1_taxonomy.png", 3000)
    render_svg(f2, figures / "figure_2_workflow.png", 3000)
    figure_validation = validate_figure_text(
        facts, f1, f2, variant=args.figure_variant
    )
    (output / "analysis" / "figure_validation.json").write_text(
        json.dumps(figure_validation, indent=2) + "\n", encoding="utf-8"
    )
    source_manifest = {
        "template": str(args.template),
        "template_sha256": sha256(args.template),
        "taxonomy_dir": str(taxonomy_dir),
        "figure_variant": args.figure_variant,
        "facts_sha256": sha256(output / "analysis" / "manuscript_fact_table.json"),
        "figure_contract_sha256": sha256(output / "analysis" / "figure_fact_contract.json"),
        "deterministic_figures": {
            "taxonomy_svg_sha256": sha256(f1),
            "taxonomy_png_sha256": sha256(figures / "figure_1_taxonomy.png"),
            "workflow_svg_sha256": sha256(f2),
            "workflow_png_sha256": sha256(figures / "figure_2_workflow.png"),
        },
    }
    (output / "source_manifest.json").write_text(
        json.dumps(source_manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    content_path = args.content or (output / "manuscript_content.json")
    if content_path.exists():
        content = read_json(content_path)
        write_manuscript_markdown(content, output)
    if args.build_docx:
        if not content_path.exists():
            raise SystemExit(f"Manuscript content not found: {content_path}")
        final_docx = build_docx(args.template, output, figures, content)
        print(final_docx)
    else:
        print(json.dumps({"facts": facts["routes"], "output": str(output)}, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
