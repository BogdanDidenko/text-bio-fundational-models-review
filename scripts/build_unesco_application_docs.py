from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


OUT_DIR = Path("data/unesco_application")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(10)
        sr = sp.add_run(subtitle)
        sr.font.name = "Calibri"
        sr.font.size = Pt(10)
        sr.font.color.rgb = RGBColor(85, 85, 85)


def set_common_styles(doc, compact=False):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75 if compact else 1.0)
    section.bottom_margin = Inches(0.75 if compact else 1.0)
    section.left_margin = Inches(0.85 if compact else 1.0)
    section.right_margin = Inches(0.85 if compact else 1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.6 if compact else 11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(5 if compact else 8)
    normal.paragraph_format.line_spacing = 1.08 if compact else 1.2

    for style_name, size, before, after in (
        ("Heading 1", 16, 14, 6),
        ("Heading 2", 13, 10, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def build_motivation_letter():
    doc = Document()
    set_common_styles(doc, compact=True)
    add_title(
        doc,
        "Motivation Letter",
        "UNESCO and AI-Lab/CBPF Remote Access to Artificial Intelligence Infrastructure",
    )

    paragraphs = [
        "I am applying for remote access to the UNESCO and AI-Lab/CBPF artificial intelligence infrastructure to advance my PhD research at Lviv Polytechnic National University on a practical question that I repeatedly encounter in bioinformatics: why do classical machine learning methods so often outperform, match, or remain preferable to deep learning and foundation-model approaches on biological tasks?",
        "My professional background is in applied machine learning and language technologies. At WebSpellChecker LLC I have worked on custom transformer-based architectures for grammatical error correction, data preprocessing, model training, and reproducible training pipelines. In parallel, my doctoral work focuses on integrative causal-oriented analysis of molecular biological markers for prioritizing anti-aging interventions. This combination has made the gap between modern AI infrastructure and biological data practice especially visible to me. Many biological datasets are small, noisy, heterogeneous, batch-affected, and expensive to validate. In these conditions, the best model is not always the largest model, and stronger benchmarking is needed before complex architectures are treated as the default.",
        "The proposed project, \"Understanding Why Classical ML Outperforms Deep LLMs in Bioinformatics,\" is motivated by this tension. I want to build a transparent benchmark and analysis workflow that compares classical ML methods, deep neural networks, and bio-foundational models across representative bioinformatics tasks. The goal is not to reject deep learning, but to identify when it provides genuine value, when simpler models are more robust, and which measurable properties of biological data explain the difference.",
        "I have already made progress toward this direction by assembling and screening literature on biological foundation models, preserving decision logs, building a model catalogue, and experimenting with document-processing workflows for full-text review. The next step is to connect this evidence base to executable benchmarks where model behavior can be tested under matched data splits, metrics, compute budgets, and reproducibility controls.",
        "Access to remote AI infrastructure is important because the project requires controlled experiments across multiple model families, datasets, random seeds, and preprocessing variants. On local resources, it is difficult to separate scientific conclusions from compute limitations. With stable GPU and CPU infrastructure, I can run systematic comparisons, keep reproducible logs, test sensitivity to data size and feature structure, and document negative as well as positive results. This is essential for a fair comparison between classical and deep approaches.",
        "My expected output is an open, reproducible benchmarking package with clear experiment logs, model-selection guidance, and a data-centric framework for deciding when bioinformatics problems benefit from foundation models. I am particularly motivated to make this work useful for researchers who do not have unlimited compute budgets and who need reliable, interpretable methods for biological discovery. The UNESCO/AI-Lab infrastructure would directly support this objective by giving the project the computational depth required to move from anecdotal comparisons to systematic evidence.",
    ]
    for text in paragraphs:
        add_body_paragraph(doc, text)

    out = OUT_DIR / "Bogdan_Didenko_Motivation_Letter_UNESCO_AI_Lab.docx"
    doc.save(out)
    return out


def build_publications_list():
    doc = Document()
    set_common_styles(doc, compact=False)
    add_title(
        doc,
        "Relevant Publications",
        "Bogdan Didenko | Google Scholar: https://scholar.google.com/citations?user=Kj-mmWgAAAAJ&hl=en",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run("Profile summary: ").bold = True
    p.add_run("WebSpellChecker LLC; Lviv Polytechnic National University. Google Scholar reports 31 citations as of 9 July 2026.")

    publications = [
        (
            "2026",
            "Data-Efficient Adaptation of Multilingual LLMs to Ukrainian",
            "Y. Paniv, B. Didenko, M. Haltiuk, V. Humennyy, A. Kravchenko, R. Kyslyi, et al.",
            "Proceedings of the Fifth Ukrainian Natural Language Processing Conference",
            "",
        ),
        (
            "2023",
            "RedPenNet for Grammatical Error Correction: Outputs to Tokens, Attentions to Spans",
            "B. Didenko, A. Sameliuk",
            "Proceedings of the Second Ukrainian Natural Language Processing Workshop",
            "10 citations",
        ),
        (
            "2019",
            "Multi-headed Architecture Based on BERT for Grammatical Errors Correction",
            "B. Didenko, J. Shaptala",
            "Proceedings of the Fourteenth Workshop on Innovative Use of NLP for Building Educational Applications",
            "21 citations",
        ),
    ]

    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    set_table_borders(table)
    headers = ["Year", "Publication", "Authors", "Venue", "Citations"]
    widths = [0.6, 2.25, 1.55, 1.65, 0.75]
    for idx, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = Inches(width)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)

    for row_data in publications:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row_data, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(value)
            run.font.size = Pt(9)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.add_run("Relevance to the application: ").bold = True
    note.add_run("these works demonstrate experience with transformer architectures, multilingual LLM adaptation, reproducible NLP benchmarking, and applied model development, which directly supports the proposed comparison of classical ML, deep learning, and foundation-model approaches.")

    out = OUT_DIR / "Bogdan_Didenko_Relevant_Publications_UNESCO_AI_Lab.docx"
    doc.save(out)
    return out


def pdf_styles():
    base = getSampleStyleSheet()
    title = ParagraphStyle(
        "AppTitle",
        parent=base["Title"],
        fontName="Helvetica-Bold",
        fontSize=21,
        leading=24,
        textColor=colors.HexColor("#0B2545"),
        alignment=TA_LEFT,
        spaceAfter=4,
    )
    subtitle = ParagraphStyle(
        "Subtitle",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=11.5,
        textColor=colors.HexColor("#555555"),
        spaceAfter=11,
    )
    body = ParagraphStyle(
        "BodyJustified",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=9.9,
        leading=11.8,
        alignment=TA_LEFT,
        spaceAfter=6,
    )
    small = ParagraphStyle(
        "Small",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=8.2,
        leading=10,
        alignment=TA_LEFT,
        spaceAfter=0,
    )
    small_bold = ParagraphStyle(
        "SmallBold",
        parent=small,
        fontName="Helvetica-Bold",
    )
    return title, subtitle, body, small, small_bold


def build_motivation_pdf():
    title, subtitle, body, _, _ = pdf_styles()
    out = OUT_DIR / "Bogdan_Didenko_Motivation_Letter_UNESCO_AI_Lab.pdf"
    doc = SimpleDocTemplate(
        str(out),
        pagesize=letter,
        rightMargin=0.78 * inch,
        leftMargin=0.78 * inch,
        topMargin=0.68 * inch,
        bottomMargin=0.68 * inch,
    )
    story = [
        Paragraph("Motivation Letter", title),
        Paragraph("UNESCO and AI-Lab/CBPF Remote Access to Artificial Intelligence Infrastructure", subtitle),
    ]
    paragraphs = [
        "I am applying for remote access to the UNESCO and AI-Lab/CBPF artificial intelligence infrastructure to advance my PhD research at Lviv Polytechnic National University on a practical question that I repeatedly encounter in bioinformatics: why do classical machine learning methods so often outperform, match, or remain preferable to deep learning and foundation-model approaches on biological tasks?",
        "My professional background is in applied machine learning and language technologies. At WebSpellChecker LLC I have worked on custom transformer-based architectures for grammatical error correction, data preprocessing, model training, and reproducible training pipelines. In parallel, my doctoral work focuses on integrative causal-oriented analysis of molecular biological markers for prioritizing anti-aging interventions. This combination has made the gap between modern AI infrastructure and biological data practice especially visible to me. Many biological datasets are small, noisy, heterogeneous, batch-affected, and expensive to validate. In these conditions, the best model is not always the largest model, and stronger benchmarking is needed before complex architectures are treated as the default.",
        "The proposed project, \"Understanding Why Classical ML Outperforms Deep LLMs in Bioinformatics,\" is motivated by this tension. I want to build a transparent benchmark and analysis workflow that compares classical ML methods, deep neural networks, and bio-foundational models across representative bioinformatics tasks. The goal is not to reject deep learning, but to identify when it provides genuine value, when simpler models are more robust, and which measurable properties of biological data explain the difference.",
        "I have already made progress toward this direction by assembling and screening literature on biological foundation models, preserving decision logs, building a model catalogue, and experimenting with document-processing workflows for full-text review. The next step is to connect this evidence base to executable benchmarks where model behavior can be tested under matched data splits, metrics, compute budgets, and reproducibility controls.",
        "Access to remote AI infrastructure is important because the project requires controlled experiments across multiple model families, datasets, random seeds, and preprocessing variants. On local resources, it is difficult to separate scientific conclusions from compute limitations. With stable GPU and CPU infrastructure, I can run systematic comparisons, keep reproducible logs, test sensitivity to data size and feature structure, and document negative as well as positive results. This is essential for a fair comparison between classical and deep approaches.",
        "My expected output is an open, reproducible benchmarking package with clear experiment logs, model-selection guidance, and a data-centric framework for deciding when bioinformatics problems benefit from foundation models. I am particularly motivated to make this work useful for researchers who do not have unlimited compute budgets and who need reliable, interpretable methods for biological discovery. The UNESCO/AI-Lab infrastructure would directly support this objective by giving the project the computational depth required to move from anecdotal comparisons to systematic evidence.",
    ]
    for text in paragraphs:
        story.append(Paragraph(text, body))
    doc.build(story)
    return out


def build_publications_pdf():
    title, subtitle, body, small, small_bold = pdf_styles()
    out = OUT_DIR / "Bogdan_Didenko_Relevant_Publications_UNESCO_AI_Lab.pdf"
    doc = SimpleDocTemplate(
        str(out),
        pagesize=letter,
        rightMargin=0.72 * inch,
        leftMargin=0.72 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
    )
    story = [
        Paragraph("Relevant Publications", title),
        Paragraph("Bogdan Didenko | Google Scholar: https://scholar.google.com/citations?user=Kj-mmWgAAAAJ&amp;hl=en", subtitle),
        Paragraph("<b>Profile summary:</b> WebSpellChecker LLC; Lviv Polytechnic National University. Google Scholar reports 31 citations as of 9 July 2026.", body),
        Spacer(1, 4),
    ]
    headers = ["Year", "Publication", "Authors", "Venue", "Citations"]
    rows = [[Paragraph(h, small_bold) for h in headers]]
    publications = [
        ["2026", "Data-Efficient Adaptation of Multilingual LLMs to Ukrainian", "Y. Paniv, B. Didenko, M. Haltiuk, V. Humennyy, A. Kravchenko, R. Kyslyi, et al.", "Proceedings of the Fifth Ukrainian Natural Language Processing Conference", ""],
        ["2023", "RedPenNet for Grammatical Error Correction: Outputs to Tokens, Attentions to Spans", "B. Didenko, A. Sameliuk", "Proceedings of the Second Ukrainian Natural Language Processing Workshop", "10 citations"],
        ["2019", "Multi-headed Architecture Based on BERT for Grammatical Errors Correction", "B. Didenko, J. Shaptala", "Proceedings of the Fourteenth Workshop on Innovative Use of NLP for Building Educational Applications", "21 citations"],
    ]
    for pub in publications:
        rows.append([Paragraph(item or " ", small) for item in pub])
    table = Table(rows, colWidths=[0.45 * inch, 2.05 * inch, 1.35 * inch, 1.75 * inch, 0.65 * inch])
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DADCE0")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(table)
    story.append(Spacer(1, 10))
    story.append(Paragraph("<b>Relevance to the application:</b> these works demonstrate experience with transformer architectures, multilingual LLM adaptation, reproducible NLP benchmarking, and applied model development, which directly supports the proposed comparison of classical ML, deep learning, and foundation-model approaches.", body))
    doc.build(story)
    return out


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print(build_motivation_letter())
    print(build_publications_list())
    print(build_motivation_pdf())
    print(build_publications_pdf())
